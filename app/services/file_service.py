"""
File Service – upload handling, validation, and structured text extraction
from PDF / DOCX files.

Features:
    * Validates file type, MIME type, and size (configurable, default 5 MB).
    * Extracts raw text while preserving section structure (headings, bullets).
    * Returns a ``ResumeDocument`` dataclass with text, sections, page count
      and file type.
    * Optionally redacts PII (emails, phone numbers, SSNs) when the
      ``PII_REDACT`` environment variable is truthy.
    * Raises ``FileParseError`` for corrupt or unreadable files.
"""

from __future__ import annotations

import logging
import mimetypes
import os
import re
import uuid
from dataclasses import dataclass, field
from typing import Dict, List, Set, Tuple

from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename

logger = logging.getLogger(__name__)

# ── optional PII redaction flag ──────────────────────────────────────
_PII_REDACT: bool = os.environ.get("PII_REDACT", "").strip().lower() in {
    "1", "true", "yes", "on",
}

# ── MIME allow-list ──────────────────────────────────────────────────
_MIME_WHITELIST: Dict[str, Set[str]] = {
    "pdf": {
        "application/pdf",
    },
    "docx": {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    },
}


# =====================================================================
# Custom exceptions
# =====================================================================

class FileParseError(Exception):
    """Raised when a file cannot be parsed (corrupt, unreadable, etc.)."""

    def __init__(self, message: str, filepath: str = "", cause: Exception | None = None) -> None:
        self.filepath = filepath
        self.cause = cause
        super().__init__(message)


# =====================================================================
# ResumeDocument data-class
# =====================================================================

@dataclass
class ResumeDocument:
    """Structured result of file extraction.

    Attributes:
        raw_text:   Full text extracted from the document.
        sections:   Mapping of detected section headings to their body text.
        page_count: Number of pages (1 for DOCX).
        file_type:  Normalised extension (``"pdf"`` or ``"docx"``).
    """

    raw_text: str = ""
    sections: Dict[str, str] = field(default_factory=dict)
    page_count: int = 0
    file_type: str = ""


# =====================================================================
# PII helpers (private)
# =====================================================================

_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")
_PHONE_RE = re.compile(
    r"(?:\+?\d{1,3}[\s\-.]?)?"          # optional country code
    r"(?:\(?\d{2,4}\)?[\s\-.]?)?"        # optional area code
    r"\d{3,4}[\s\-.]?\d{3,4}"           # main number
)
_SSN_RE = re.compile(r"\b\d{3}-\d{2}-\d{4}\b")


def _redact(text: str) -> str:
    """Replace common PII tokens with placeholders."""
    text = _EMAIL_RE.sub("[EMAIL REDACTED]", text)
    text = _SSN_RE.sub("[SSN REDACTED]", text)
    text = _PHONE_RE.sub("[PHONE REDACTED]", text)
    return text


# =====================================================================
# Section detection (private)
# =====================================================================

_SECTION_HEADINGS = re.compile(
    r"^(EDUCATION|EXPERIENCE|WORK\s*EXPERIENCE|SKILLS|PROJECTS|"
    r"CERTIFICATIONS|CERTIFICATES|SUMMARY|OBJECTIVE|PROFILE|"
    r"PROFESSIONAL\s*SUMMARY|TECHNICAL\s*SKILLS|ACHIEVEMENTS|"
    r"AWARDS|PUBLICATIONS|REFERENCES|INTERESTS|HOBBIES|"
    r"LANGUAGES|VOLUNTEER|ACTIVITIES)\b",
    re.IGNORECASE | re.MULTILINE,
)


def _split_sections(text: str) -> Dict[str, str]:
    """Detect common resume headings and split the text into sections.

    Returns a dict mapping heading names (title-cased) to the body text
    beneath them.  A special ``"header"`` key captures any text above
    the first detected heading.
    """
    matches = list(_SECTION_HEADINGS.finditer(text))
    if not matches:
        return {"header": text.strip()}

    sections: Dict[str, str] = {}

    # Text before first heading
    preamble = text[: matches[0].start()].strip()
    if preamble:
        sections["Header"] = preamble

    for idx, match in enumerate(matches):
        heading = " ".join(match.group().split()).title()
        start = match.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        body = text[start:end].strip()
        sections[heading] = body

    return sections


# =====================================================================
# FileService
# =====================================================================

class FileService:
    """Handles secure file uploads, validation, and structured text
    extraction from PDF and DOCX documents.

    Class Attributes:
        ALLOWED_EXTENSIONS: Extensions accepted for upload.
        MAX_FILE_SIZE:      Maximum upload size in bytes (default 5 MB).
    """

    ALLOWED_EXTENSIONS: Set[str] = {"pdf", "docx"}
    MAX_FILE_SIZE: int = 5 * 1024 * 1024  # 5 MB

    # ------------------------------------------------------------------
    # Validation
    # ------------------------------------------------------------------

    @staticmethod
    def is_allowed(filename: str, allowed: Set[str] | None = None) -> bool:
        """Check whether *filename* has a permitted extension.

        Args:
            filename: Original upload filename.
            allowed:  Set of extensions to allow (defaults to
                      ``ALLOWED_EXTENSIONS``).

        Returns:
            ``True`` if the extension is accepted.
        """
        allowed = allowed or FileService.ALLOWED_EXTENSIONS
        return "." in filename and filename.rsplit(".", 1)[1].lower() in allowed

    @staticmethod
    def validate(file_obj: FileStorage) -> None:
        """Validate an uploaded file for type, MIME, and size.

        Args:
            file_obj: Werkzeug ``FileStorage`` from the request.

        Raises:
            FileParseError: If any validation check fails.
        """
        filename = file_obj.filename or ""
        if not filename:
            raise FileParseError("No filename provided")

        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext not in FileService.ALLOWED_EXTENSIONS:
            raise FileParseError(
                f"Unsupported file type '.{ext}'. Allowed: "
                f"{', '.join(sorted(FileService.ALLOWED_EXTENSIONS))}",
                filepath=filename,
            )

        # MIME check – use the content_type reported by the browser and
        # fall back to mimetypes for extra safety.
        reported_mime = (file_obj.content_type or "").lower()
        guessed_mime, _ = mimetypes.guess_type(filename)
        guessed_mime = (guessed_mime or "").lower()
        allowed_mimes = _MIME_WHITELIST.get(ext, set())
        if allowed_mimes and reported_mime not in allowed_mimes and guessed_mime not in allowed_mimes:
            raise FileParseError(
                f"Invalid MIME type '{reported_mime}' for .{ext} file",
                filepath=filename,
            )

        # Size check – read current position, seek to end, measure, seek back.
        pos = file_obj.stream.tell()
        file_obj.stream.seek(0, os.SEEK_END)
        size = file_obj.stream.tell()
        file_obj.stream.seek(pos)

        if size > FileService.MAX_FILE_SIZE:
            max_mb = FileService.MAX_FILE_SIZE / (1024 * 1024)
            raise FileParseError(
                f"File too large ({size / (1024 * 1024):.1f} MB). "
                f"Maximum allowed: {max_mb:.0f} MB",
                filepath=filename,
            )

    # ------------------------------------------------------------------
    # Upload
    # ------------------------------------------------------------------

    @staticmethod
    def save_upload(file_obj: FileStorage, upload_folder: str) -> str:
        """Validate and save an uploaded file with a unique name.

        Args:
            file_obj:      Werkzeug ``FileStorage``.
            upload_folder: Directory to save into.

        Returns:
            Absolute path of the saved file.

        Raises:
            FileParseError: If validation or saving fails.
        """
        FileService.validate(file_obj)
        os.makedirs(upload_folder, exist_ok=True)
        try:
            safe_name = secure_filename(file_obj.filename or "upload")
            unique_name = f"{uuid.uuid4().hex}_{safe_name}"
            dest = os.path.join(upload_folder, unique_name)
            file_obj.stream.seek(0)
            file_obj.save(dest)
            logger.info("Saved upload: %s", unique_name)
            return dest
        except Exception as exc:
            raise FileParseError(
                f"Failed to save file: {exc}",
                filepath=file_obj.filename or "",
                cause=exc,
            ) from exc

    # ------------------------------------------------------------------
    # Text extraction
    # ------------------------------------------------------------------

    @staticmethod
    def extract(filepath: str) -> ResumeDocument:
        """Extract text and structure from a PDF or DOCX file.

        Args:
            filepath: Path to the file on disk.

        Returns:
            A populated ``ResumeDocument``.

        Raises:
            FileParseError: If the file is corrupt or unreadable.
        """
        if not os.path.isfile(filepath):
            raise FileParseError(f"File not found: {filepath}", filepath=filepath)

        ext = os.path.splitext(filepath)[1].lower().lstrip(".")
        try:
            if ext == "pdf":
                raw_text, page_count = FileService._read_pdf(filepath)
            elif ext in ("docx", "doc"):
                raw_text, page_count = FileService._read_docx(filepath)
            else:
                raise FileParseError(
                    f"Unsupported file type: .{ext}", filepath=filepath
                )
        except FileParseError:
            raise
        except Exception as exc:
            raise FileParseError(
                f"Failed to read {ext.upper()} file: {exc}",
                filepath=filepath,
                cause=exc,
            ) from exc

        if not raw_text.strip():
            raise FileParseError(
                "No text could be extracted (file may be scanned/image-only)",
                filepath=filepath,
            )

        if _PII_REDACT:
            raw_text = _redact(raw_text)

        sections = _split_sections(raw_text)

        return ResumeDocument(
            raw_text=raw_text,
            sections=sections,
            page_count=page_count,
            file_type=ext,
        )

    @staticmethod
    def extract_fast(filepath: str) -> ResumeDocument:
        """Fast text extraction WITHOUT section detection (optimized for batch).
        
        Skips the section detection step for better performance in batch operations.
        Use this when you only need raw text, not structured sections.

        Args:
            filepath: Path to the file on disk.

        Returns:
            A ResumeDocument with raw_text, page_count, and file_type (sections empty).

        Raises:
            FileParseError: If the file is corrupt or unreadable.
        """
        if not os.path.isfile(filepath):
            raise FileParseError(f"File not found: {filepath}", filepath=filepath)

        ext = os.path.splitext(filepath)[1].lower().lstrip(".")
        try:
            if ext == "pdf":
                raw_text, page_count = FileService._read_pdf(filepath)
            elif ext in ("docx", "doc"):
                raw_text, page_count = FileService._read_docx(filepath)
            else:
                raise FileParseError(
                    f"Unsupported file type: .{ext}", filepath=filepath
                )
        except FileParseError:
            raise
        except Exception as exc:
            raise FileParseError(
                f"Failed to read {ext.upper()} file: {exc}",
                filepath=filepath,
                cause=exc,
            ) from exc

        if not raw_text.strip():
            raise FileParseError(
                "No text could be extracted (file may be scanned/image-only)",
                filepath=filepath,
            )

        if _PII_REDACT:
            raw_text = _redact(raw_text)

        # Skip section detection for performance
        return ResumeDocument(
            raw_text=raw_text,
            sections={},  # Empty for fast extraction
            page_count=page_count,
            file_type=ext,
        )

    # ------------------------------------------------------------------
    # PDF (private)
    # ------------------------------------------------------------------

    @staticmethod
    def _read_pdf(filepath: str) -> Tuple[str, int]:
        """Extract text from a PDF using *pdfplumber* with fallback to pymupdf.

        Returns:
            ``(text, page_count)``

        Raises:
            FileParseError: If extraction fails.
        """
        try:
            import pdfplumber  # type: ignore[import-untyped]
        except ImportError as exc:
            raise FileParseError(
                "pdfplumber is required for PDF extraction. "
                "Install it with: pip install pdfplumber",
                filepath=filepath,
                cause=exc,
            ) from exc

        try:
            pages: List[str] = []
            with pdfplumber.open(filepath) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    try:
                        text = page.extract_text()
                        if text:
                            pages.append(text)
                    except Exception:
                        # Skip problematic pages silently
                        pass
            
            if pages:
                return "\n\n".join(pages), page_count
            
            # If pdfplumber extracted nothing, try fallback
            logger.warning("No pages extracted with pdfplumber, attempting fallback")
            return FileService._read_pdf_fallback(filepath)
            
        except FileParseError:
            raise
        except Exception as exc:
            # Any other error, try fallback
            logger.warning(f"pdfplumber extraction failed: {exc}, trying fallback")
            try:
                return FileService._read_pdf_fallback(filepath)
            except FileParseError:
                raise
            except Exception:
                raise FileParseError(
                    f"PDF extraction failed (pdfplumber, then fallback): {exc}",
                    filepath=filepath,
                    cause=exc,
                ) from exc

    @staticmethod
    def _read_pdf_fallback(filepath: str) -> Tuple[str, int]:
        """Fallback PDF extraction using pymupdf (fitz).
        
        This is a simpler, more robust extraction method for PDFs that
        fail with pdfplumber (e.g., due to font encoding issues).
        
        Returns:
            ``(text, page_count)``
        """
        try:
            import fitz  # type: ignore[import-untyped]
        except ImportError:
            raise FileParseError(
                "Could not extract PDF with fallback method. "
                "Try installing: pip install pymupdf",
                filepath=filepath,
            )

        try:
            pages: List[str] = []
            with fitz.open(filepath) as pdf:
                page_count = len(pdf)
                for page_num in range(page_count):
                    page = pdf[page_num]
                    text = page.get_text()
                    if text:
                        pages.append(text.strip())
            
            if not pages:
                raise FileParseError(
                    "No text could be extracted (file may be scanned/image-only)",
                    filepath=filepath,
                )
            
            return "\n\n".join(pages), page_count
        except FileParseError:
            raise
        except Exception as exc:
            raise FileParseError(
                f"Fallback PDF extraction failed: {exc}",
                filepath=filepath,
                cause=exc,
            ) from exc

    @staticmethod
    def _read_docx(filepath: str) -> Tuple[str, int]:
        """Extract text from a DOCX file preserving headings and bullets.

        Headings are emitted in UPPERCASE so that ``_split_sections``
        can detect them.  Bullet / list paragraphs are prefixed with
        ``"• "``.

        Returns:
            ``(text, 1)`` – DOCX files report a single logical page.

        Raises:
            FileParseError: If python-docx cannot open the file.
        """
        try:
            from docx import Document  # type: ignore[import-untyped]
        except ImportError as exc:
            raise FileParseError(
                "python-docx is required for DOCX extraction.",
                filepath=filepath,
                cause=exc,
            ) from exc

        try:
            doc = Document(filepath)
        except Exception as exc:
            raise FileParseError(
                f"Corrupt or unreadable DOCX: {exc}",
                filepath=filepath,
                cause=exc,
            ) from exc

        parts: List[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name or "").lower() if para.style else ""
            if "heading" in style_name:
                parts.append(f"\n{text.upper()}\n")
            elif "list" in style_name or "bullet" in style_name:
                parts.append(f"• {text}")
            else:
                parts.append(text)

        # Also capture table content
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts), 1

    # ------------------------------------------------------------------
    # Utilities
    # ------------------------------------------------------------------

    @staticmethod
    def delete_file(filepath: str) -> bool:
        """Delete a file from disk.

        Args:
            filepath: Absolute path.

        Returns:
            ``True`` if the file was removed, ``False`` otherwise.
        """
        try:
            if os.path.isfile(filepath):
                os.remove(filepath)
                return True
        except OSError as exc:
            logger.error("Delete failed: %s", exc)
        return False
